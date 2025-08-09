import os
import json
from collections import defaultdict
from dataclasses import dataclass
from typing import List, Dict, Any, Tuple

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ---------------------------------------------------------------------------------
# Step 0 (Why): Introduce additional helpers for ToC, page fields, numbering, and paths
# - The reference report has a title page, ToC, numbered sections, and rich tables
# - These helpers enable Word fields (TOC, PAGE) and standardize asset paths
# ---------------------------------------------------------------------------------
# Step 0 (Done): Helpers added

OUTPUT_DIR = os.path.join("output")
DATA_DIR = os.path.join(OUTPUT_DIR, "data")
DRILLDOWN_DIR = os.path.join(OUTPUT_DIR, "graphs", "drilldown")
SUMMARY_DIR = os.path.join(OUTPUT_DIR, "graphs", "summary")

INPUT_QUALITATIVE_JSON = os.path.join(DATA_DIR, 'qualitative_analysis.json')
INPUT_QUANTITATIVE_JSON = os.path.join(DATA_DIR, 'quantitative_results.json')
OUTPUT_DOCX = os.path.join(OUTPUT_DIR, 'Final_Analysis_Report.docx')

# Optional logo if available locally
LOGO_PATH = os.path.join("assets", "logo.png")


def sanitize_component(text: str) -> str:
    return ''.join(ch for ch in str(text) if ch.isalnum() or ch in ('-', '_')).replace(' ', '')


def drilldown_path(analysis_id: str) -> str:
    return os.path.join(DRILLDOWN_DIR, f"{analysis_id}.png")


def summary_path(language: str, dimension: str) -> str:
    return os.path.join(SUMMARY_DIR, f"{sanitize_component(language)}-{sanitize_component(dimension)}_summary.png")


def _add_field_code(run, field_text: str) -> None:
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = field_text
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    r_element = run._r
    r_element.append(fldChar1)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar3)


# ---------------------------------------------------------------------------------
# Step 1 (Why): Add Title and ToC writers to mirror the reference
# - Title page: logo (if present), title, author, contact
# - ToC: True Word field, so users can update ToC in Word
# ---------------------------------------------------------------------------------
# Step 1 (Done): TitlePageWriterAgent and TableOfContentsWriterAgent added

class TitlePageWriterAgent:
    def write(self, doc: Document, title: str, author: str, contact_email: str, contact_phone: str):
        # Centered title area
        if os.path.exists(LOGO_PATH):
            p_logo = doc.add_paragraph()
            p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_logo = p_logo.add_run()
            run_logo.add_picture(LOGO_PATH, width=Inches(2.0))
        p_title = doc.add_paragraph()
        p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_t = p_title.add_run(title)
        run_t.bold = True
        run_t.font.size = doc.styles['Title'].font.size
        # Author and contact
        p_author = doc.add_paragraph()
        p_author.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_author.add_run(author)
        p_contact = doc.add_paragraph()
        p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_contact.add_run(f"Email: {contact_email} | Mobile: {contact_phone}")
        doc.add_page_break()


class TableOfContentsWriterAgent:
    def write(self, doc: Document, heading: str = 'Table of Contents'):
        doc.add_heading(heading, level=1)
        p = doc.add_paragraph()
        run = p.add_run()
        _add_field_code(run, 'TOC \\o "1-3" \\h \\z \\u')
        doc.add_page_break()


# ---------------------------------------------------------------------------------
# Step 2 (Why): Upgrade Summary tables to match reference-style matrices per dimension
# - Build tables for Relevance, Completeness, Correctness with model 5-counts and winner
# - Uses quantitative_entries contract
# ---------------------------------------------------------------------------------
# Step 2 (Done): TabularDataWriterAgent extended

MODEL_DISPLAY_ORDER = [
    'OpenAI o4-mini-high',
    'Gemini 2.5 Pro',
    'Claude Opus 4'
]
MODEL_SHORT_TO_DISPLAY = {
    'Ollama': 'OpenAI o4-mini-high',
    'Gemini': 'Gemini 2.5 Pro',
    'Claude': 'Claude Opus 4',
    'Grok': 'xAI Grok 4'
}


def _score5_for_model_row(models: Dict[str, Any], model_short: str) -> int:
    data = models.get(model_short, {})
    return int(data.get('score_5', 0))


class TabularDataWriterAgent:
    def write(self, doc: Document, quantitative_entries: List[Dict[str, Any]]):
        if not quantitative_entries:
            return
        # Group rows by dimension
        by_dimension: Dict[str, List[Dict[str, Any]]] = defaultdict(list)
        for e in quantitative_entries:
            by_dimension[e['dimension']].append(e)
        # For each dimension, emit a summary table like the reference
        for dimension in sorted(by_dimension.keys()):
            doc.add_heading(f"{dimension} — Summary Table", level=1)
            table = doc.add_table(rows=1, cols=6)
            hdr = table.rows[0].cells
            hdr[0].text = 'Language'
            hdr[1].text = 'Industry'
            hdr[2].text = 'OpenAI o4-mini-high (5s)'
            hdr[3].text = 'Gemini 2.5 Pro (5s)'
            hdr[4].text = 'Claude Opus 4 (5s)'
            hdr[5].text = 'Winner'
            # Populate rows, sorted by language then domain
            for row in sorted(by_dimension[dimension], key=lambda x: (x['language'], x['domain'])):
                models = row.get('models', {})
                # Models in JSON are keyed by short names (e.g., 'Gemini', 'Claude', 'Ollama', 'Grok')
                # Map to display
                scores_display = {
                    'OpenAI o4-mini-high': _score5_for_model_row(models, 'Ollama'),
                    'Gemini 2.5 Pro': _score5_for_model_row(models, 'Gemini'),
                    'Claude Opus 4': _score5_for_model_row(models, 'Claude')
                }
                w = row.get('winner', {}).get('winner', '')
                winner_display = MODEL_SHORT_TO_DISPLAY.get(w, w)
                cells = table.add_row().cells
                cells[0].text = str(row.get('language', ''))
                cells[1].text = str(row.get('domain', ''))
                cells[2].text = str(scores_display['OpenAI o4-mini-high'])
                cells[3].text = str(scores_display['Gemini 2.5 Pro'])
                cells[4].text = str(scores_display['Claude Opus 4'])
                cells[5].text = str(winner_display)
            doc.add_page_break()


# ---------------------------------------------------------------------------------
# Step 3 (Why): Numbered qualitative sections mirroring 4.x structure from reference
# - Insert summary image once per (Language, Dimension), then per industry entry
# - Provide numbering counters for Dimension → Language → Industry as 4.i, 4.i.j, 4.i.j.k
# ---------------------------------------------------------------------------------
# Step 3 (Done): QualitativeAnalysisWriterAgent updated with numbering

class QualitativeAnalysisWriterAgent:
    def write(self, doc: Document, qualitative_entries: List[Dict[str, Any]]):
        if not qualitative_entries:
            return
        by_dim_lang: Dict[str, Dict[str, List[Dict[str, Any]]]] = defaultdict(lambda: defaultdict(list))
        for e in qualitative_entries:
            by_dim_lang[e['dimension']][e['language']].append(e)
        base_section = 4
        dim_index = 0
        for dimension in sorted(by_dim_lang.keys()):
            dim_index += 1
            doc.add_heading(f"{base_section}.{dim_index} {dimension}", level=1)
            for language in sorted(by_dim_lang[dimension].keys()):
                # Language summary heading
                lang_index = 1
                doc.add_heading(f"{base_section}.{dim_index}.{lang_index} {language} Language Summary", level=2)
                # Insert summary grid image for this (language, dimension)
                sp = summary_path(language, dimension)
                if os.path.exists(sp):
                    doc.add_picture(sp, width=Inches(6))
                # Per-industry subsections
                industry_index = 0
                for e in sorted(by_dim_lang[dimension][language], key=lambda x: x['domain']):
                    industry_index += 1
                    doc.add_heading(f"{base_section}.{dim_index}.{lang_index}.{industry_index} {e['domain']} ({e['language']})", level=3)
                    dp = drilldown_path(e['analysis_id'])
                    if os.path.exists(dp):
                        doc.add_picture(dp, width=Inches(5.5))
                    if e.get('winner_text'):
                        doc.add_paragraph(e['winner_text'])
                    if e.get('client_performance_text'):
                        doc.add_paragraph(e['client_performance_text'])
            doc.add_page_break()


# ---------------------------------------------------------------------------------
# Step 4 (Why): Add Conclusion/Recommendations placeholders to match reference end sections
# - Keep existing sections and ensure consistent heading levels
# ---------------------------------------------------------------------------------
# Step 4 (Done): RecommendationWriterAgent unchanged, added heading level alignment

class RecommendationWriterAgent:
    def write(self, doc: Document):
        doc.add_heading('5. Conclusion', level=1)
        doc.add_paragraph('This section synthesizes key findings from the analyses across dimensions and languages.')
        doc.add_heading('6. Recommendations', level=1)
        doc.add_paragraph('Prioritize models that consistently achieve higher “5” counts on critical dimensions for each industry.')


# ---------------------------------------------------------------------------------
# Step 5 (Why): Orchestration upgraded to include Title, ToC, Summaries, Qualitative
# - Sequence: Title → ToC → Summary Tables → Qualitative Sections → Conclusion/Recommendations
# - Add footer page number field for polish
# ---------------------------------------------------------------------------------
# Step 5 (Done): DocumentAssemblyAgent.run upgraded

@dataclass
class DocumentAssemblyAgent:
    qualitative_entries: List[Dict[str, Any]]
    quantitative_entries: List[Dict[str, Any]]

    def _add_footer_page_numbers(self, doc: Document):
        section = doc.sections[0]
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run('Page ')
        run2 = p.add_run()
        _add_field_code(run2, 'PAGE')

    def run(self) -> str:
        doc = Document()
        # Title page
        TitlePageWriterAgent().write(
            doc,
            title='Evaluating Code Generation for Industry-Specific Software Solutions: Automated Report',
            author='By Ethara.AI',
            contact_email='suryansh@ethara.ai',
            contact_phone='+91 8595887825'
        )
        # Table of Contents (update in Word: F9)
        TableOfContentsWriterAgent().write(doc)
        # Summary tables per dimension
        TabularDataWriterAgent().write(doc, self.quantitative_entries)
        # Detailed qualitative sections with numbering
        QualitativeAnalysisWriterAgent().write(doc, self.qualitative_entries)
        # Conclusions
        RecommendationWriterAgent().write(doc)
        # Footer page numbers
        self._add_footer_page_numbers(doc)
        os.makedirs(OUTPUT_DIR, exist_ok=True)
        doc.save(OUTPUT_DOCX)
        return OUTPUT_DOCX


# ---------------------------------------------------------------------------------
# Step 6 (Why): Main entrypoint unchanged except now produces richer formatting
# ---------------------------------------------------------------------------------
# Step 6 (Done): main() keeps same behavior


def main():
    print("📄 document_assembler starting (Stage 3)")
    with open(INPUT_QUALITATIVE_JSON, 'r', encoding='utf-8') as f:
        qualitative_entries = json.load(f)
    with open(INPUT_QUANTITATIVE_JSON, 'r', encoding='utf-8') as f:
        quantitative_entries = json.load(f)

    assembler = DocumentAssemblyAgent(
        qualitative_entries=qualitative_entries,
        quantitative_entries=quantitative_entries
    )
    out = assembler.run()
    print(f"✅ Document written → {out}")
    print("🎉 document_assembler complete")


if __name__ == "__main__":
    main() 